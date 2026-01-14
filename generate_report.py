import datetime
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: 'python-docx' library is required.")
    print("Please install it using: pip install python-docx")
    exit()

def create_document():
    doc = Document()

    # --- STYLES ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title Style
    title_style = doc.styles['Title']
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)

    # Heading 1 Style
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 0) # Black, not blue

    # Heading 2 Style
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 0, 0)

    # --- 1. COVER PAGE ---
    
    # Header
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.text = "VITYARTHI"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Title Block
    for _ in range(3): doc.add_paragraph() # Spacing
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PROJECT REPORT")
    run.font.size = Pt(20)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Diabetes Prediction Web Application")
    run.font.size = Pt(24)
    run.font.bold = True
    
    for _ in range(5): doc.add_paragraph() # Spacing
    
    # Student Details
    details = [
        ("Name:", "VEDANSHIKA SINGH"),
        ("Registration No.:", "25BAI10612"),
        ("Domain:", "Healthcare, Machine Learning & Prediction"),
        ("Course:", "INTRODUCTION TO PROBLEM SOLVING AND PROGRAMMING")
    ]
    
    for label, value in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(f"{label}  ")
        run_label.bold = True
        run_value = p.add_run(value)

    doc.add_page_break()

    # --- 2. INTRODUCTION ---
    doc.add_heading('1. Introduction', level=1)
    p = doc.add_paragraph(
        "The Diabetes Prediction Web Application is a machine learning-based tool designed "
        "to assist in the early identification of diabetes. Built using Python and Streamlit, "
        "this application leverages clinical data such as Glucose levels, BMI, and Age to "
        "predict the likelihood of a patient being diabetic. The project integrates data "
        "visualization, statistical analysis, and automated reporting to provide a comprehensive "
        "tool for both general users and healthcare enthusiasts."
    )
    
    # --- 3. PROBLEM STATEMENT ---
    doc.add_heading('2. Problem Statement', level=1)
    p = doc.add_paragraph(
        "Diabetes is a prevalent chronic disease with severe long-term complications if left "
        "undiagnosed. A significant portion of the population remains unaware of their "
        "diabetic status due to irregular medical check-ups or lack of accessible testing tools. "
        "The problem this project addresses is the need for a quick, accessible, and data-driven "
        "preliminary assessment tool. By utilizing historical health data and machine learning, "
        "users can assess their risk profile instantly without immediate invasive testing, encouraging "
        "timely medical consultation."
    )

    # --- 4. FUNCTIONAL REQUIREMENTS ---
    doc.add_heading('3. Functional Requirements', level=1)
    doc.add_paragraph("The system implements the following core functional modules:")
    
    p = doc.add_paragraph("3.1 Data Input & Validation", style='List Bullet')
    doc.add_paragraph(
        "   - Users can input medical parameters (Pregnancies, Glucose, Blood Pressure, Skin Thickness, Insulin, BMI, Diabetes Pedigree Function, Age).", 
        style='Normal'
    )
    doc.add_paragraph(
        "   - The system validates inputs (e.g., ensuring non-negative values) before processing.", 
        style='Normal'
    )

    p = doc.add_paragraph("3.2 Prediction Engine", style='List Bullet')
    doc.add_paragraph(
        "   - The system loads a pre-trained Gradient Boosting Classifier model.", style='Normal'
    )
    doc.add_paragraph(
        "   - It calculates the classification (Diabetic/Non-Diabetic) and the probability score.", style='Normal'
    )

    p = doc.add_paragraph("3.3 Visualization & Reporting", style='List Bullet')
    doc.add_paragraph(
        "   - Interactive charts (Histograms, Scatter plots) allow users to explore the dataset.", style='Normal'
    )
    doc.add_paragraph(
        "   - The system generates a downloadable HTML report containing the patient's specific results and inputs.", style='Normal'
    )

    # --- 5. NON-FUNCTIONAL REQUIREMENTS ---
    doc.add_heading('4. Non-Functional Requirements', level=1)
    
    reqs = [
        ("Usability", "The interface is built with Streamlit, ensuring a clean, responsive, and intuitive user experience requiring no technical expertise."),
        ("Performance", "Predictions are generated in real-time (milliseconds) using an optimized joblib-loaded model."),
        ("Reliability", "The system handles missing data (zero values) via imputation during the training phase to ensure robust predictions."),
        ("Portability", "Being a web-based application, it can be accessed from any device with a browser and internet connection.")
    ]
    for title, desc in reqs:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    # --- 6. SYSTEM ARCHITECTURE ---
    doc.add_heading('5. System Architecture', level=1)
    doc.add_paragraph(
        "The application follows a standard Model-View-Controller (MVC) inspired architecture suited for data science applications:"
    )
    doc.add_paragraph(
        "1. Presentation Layer (Frontend): Built with Streamlit, handling user inputs and rendering charts.\n"
        "2. Logic Layer (Backend): Python scripts (app.py) handle control flow, data processing, and report generation.\n"
        "3. Model Layer: Scikit-learn Gradient Boosting model, serialized using Joblib for persistence.\n"
        "4. Data Layer: 'diabetes.csv' acts as the training source, with pandas used for dataframe manipulation."
    )

    # --- 7. DESIGN DIAGRAMS ---
    doc.add_heading('6. Design Diagrams', level=1)
    doc.add_paragraph("Note: Please refer to the attached design documents/images for visual representations.")
    
    doc.add_heading('6.1 Use Case Diagram', level=2)
    doc.add_paragraph(
        "Actors: User, System.\n"
        "Use Cases: Input Medical Data, View Prediction, Download Report, Explore Data Visualizations."
    )
    
    doc.add_heading('6.2 Workflow Diagram', level=2)
    doc.add_paragraph(
        "Start -> Load Data/Model -> User Input -> Validation -> "
        "Prediction Model -> Display Result -> (Optional) Download Report -> End."
    )

    # --- 8. DESIGN DECISIONS ---
    doc.add_heading('7. Design Decisions & Rationale', level=1)
    
    decisions = [
        ("Gradient Boosting Classifier", "Chosen over Decision Trees or Logistic Regression due to its superior performance on tabular medical data and ability to handle complex non-linear relationships."),
        ("Streamlit Framework", "Selected for its ability to rapidly convert Python data scripts into shareable web apps without requiring extensive HTML/CSS/JS knowledge."),
        ("Joblib Serialization", "Used instead of Pickle for more efficient storage of large numpy arrays typically found in Scikit-learn models.")
    ]
    
    for title, desc in decisions:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    # --- 9. IMPLEMENTATION DETAILS ---
    doc.add_heading('8. Implementation Details', level=1)
    
    doc.add_heading('8.1 Model Training (train.py)', level=2)
    doc.add_paragraph(
        "The training pipeline handles data preprocessing and model creation. Key steps include:\n"
        "- Data Cleaning: Replacing invalid zero values in columns like Glucose and BMI with NaN.\n"
        "- Imputation: Using SimpleImputer (mean strategy) to fill missing values.\n"
        "- Scaling: Applying StandardScaler to normalize feature range.\n"
        "- Algorithm: GradientBoostingClassifier is trained and saved as 'diabetes_model.joblib'."
    )

    doc.add_heading('8.2 Application Interface (app.py)', level=2)
    doc.add_paragraph(
        "The main application logic resides here:\n"
        "- Sidebar: Contains navigation for 'Home', 'Prediction', and 'Data Exploration'.\n"
        "- Prediction Page: Collects user inputs using st.number_input and calls the loaded model.\n"
        "- Report Generation: A custom function 'create_html_report' formats the result into a downloadable file."
    )

    # --- 10. SCREENSHOTS ---
    doc.add_heading('9. Screenshots / Results', level=1)
    doc.add_paragraph("[Insert Screenshot of Home Page Here]")
    doc.add_paragraph("[Insert Screenshot of Prediction Result Here]")
    doc.add_paragraph("[Insert Screenshot of Confusion Matrix Here]")
    doc.add_paragraph("The application successfully classifies inputs and displays the probability score (e.g., 'Probability of Diabetes: 65%').")

    # --- 11. TESTING APPROACH ---
    doc.add_heading('10. Testing Approach', level=1)
    doc.add_paragraph(
        "Testing was conducted in three phases as outlined in the README:"
    )
    doc.add_paragraph("1. Input Validation Testing: Verified that extreme values or negative numbers are handled or rejected.", style='List Bullet')
    doc.add_paragraph("2. Model Behavior Testing: Deleted the .joblib file and ran train.py to ensure the retraining pipeline works correctly.", style='List Bullet')
    doc.add_paragraph("3. Report Generation Testing: Verified that the downloaded HTML report correctly reflects the input data and prediction result.", style='List Bullet')

    # --- 12. CHALLENGES FACED ---
    doc.add_heading('11. Challenges Faced', level=1)
    doc.add_paragraph(
        "- Data Quality: The dataset contained zero values for physiological parameters (like Blood Pressure), which is biologically impossible. This was resolved by replacing zeros with NaN and imputing the mean.\n"
        "- Model Persistence: ensuring the Streamlit app could seamlessly load the model file created by a separate script without path errors.\n"
        "- Library Compatibility: Ensuring FPDF and Streamlit versions were compatible for report generation (switched to HTML generation for better compatibility)."
    )

    # --- 13. LEARNINGS ---
    doc.add_heading('12. Learnings & Key Takeaways', level=1)
    doc.add_paragraph(
        "- End-to-End ML Pipeline: Learned how to connect a trained backend model to a user-facing frontend.\n"
        "- Data Preprocessing Importance: Understood that model accuracy is heavily dependent on how missing values are handled.\n"
        "- Streamlit capabilities: Gained proficiency in creating interactive web layouts using Python exclusively."
    )

    # --- 14. FUTURE ENHANCEMENTS ---
    doc.add_heading('13. Future Enhancements', level=1)
    doc.add_paragraph(
        "- Database Integration: Connect a SQL database to store patient history and track trends over time.\n"
        "- Multi-Model Support: Allow users to switch between different algorithms (e.g., Random Forest vs. SVM) to compare results.\n"
        "- Mobile Optimization: Refine the UI for better usability on smaller smartphone screens."
    )

    # --- 15. REFERENCES ---
    doc.add_heading('14. References', level=1)
    doc.add_paragraph("1. Scikit-learn Documentation: https://scikit-learn.org/stable/")
    doc.add_paragraph("2. Streamlit Documentation: https://docs.streamlit.io/")
    doc.add_paragraph("3. PIMA Indians Diabetes Dataset (Kaggle).")

    # Save Document
    file_name = 'Project_Report_Vedanshika_Singh.docx'
    doc.save(file_name)
    print(f"Report generated successfully: {file_name}")

if __name__ == "__main__":
    create_document()